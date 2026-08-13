import os
import shutil
import time
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def update_page_numbers():
    report_dir = os.path.join(os.path.dirname(__file__), '..', 'report')
    target_docx = os.path.join(report_dir, 'PennyWise_Final_Project_Report.docx')
    output_docx = os.path.join(report_dir, 'PennyWise_Project_Report_Complete.docx')

    # Map of Contents Table rows to Page Numbers
    contents_page_numbers = [
        7,   # 1 Introduction
        9,   # 2 Basic Concepts
        10,  # Literature Review
        11,  # 3 Problem Statement and Requirement Specifications
        11,  # 3.1 Project Planning
        11,  # 3.2 Project Analysis (SRS)
        11,  # 3.3 System Design
        11,  # 3.3.1 Design Constraints
        12,  # 3.3.2 System Architecture / Block Diagram
        16,  # 4 Implementation
        16,  # 4.1 Methodology
        17,  # 4.2 Testing / Verification Plan
        18,  # 4.3 Result Analysis / Screenshots
        31,  # 4.4 Quality Assurance
        32,  # 5 Standard Adopted
        32,  # 5.1 Design Standards
        32,  # 5.2 Coding Standards
        32,  # 5.3 Testing Standards
        33,  # 6 Conclusion and Future Scope
        33,  # 6.1 Conclusion
        33,  # 6.2 Future Scope
        34   # References
    ]

    # Map of Figure names to Page Numbers
    figure_page_numbers = {
        'Figure 3.1': ('Figure 3.1  High-Level System Architecture', 12),
        'Figure 3.2': ('Figure 3.2  Use Case Diagram', 13),
        'Figure 3.3': ('Figure 3.3  Entity-Relationship (ER) Diagram', 14),
        'Figure 3.4': ('Figure 3.4  Data Flow / Request-Response Flow', 15),
        'Figure 4.1': ('Figure 4.1  PennyWise Dashboard', 18),
        'Figure 4.2': ('Figure 4.2  Registration / Login Interface', 19),
        'Figure 4.3': ('Figure 4.3  Expense Management Interface', 20),
        'Figure 4.4': ('Figure 4.4  Analytics and Mood Analysis', 21),
        'Figure 4.5': ('Figure 4.5  AI Behavior Suite', 22),
        'Figure 4.6': ('Figure 4.6  Should I Buy It? Advisor Result', 23),
        'Figure 4.7': ('Figure 4.7  Regret Predictor Result', 24),
        'Figure 4.8': ('Figure 4.8  Dream Savings Goals', 25),
        'Figure 4.9': ('Figure 4.9  Subscription Killer', 26),
        'Figure 4.10': ('Figure 4.10  Money Game, XP and Badges', 27),
        'Figure 4.11': ('Figure 4.11  No-Spend Calendar', 28),
        'Figure 4.12': ('Figure 4.12  Monthly Report / PDF', 29),
        'Figure 4.13': ('Figure 4.13  Profile and Deduction Management', 30),
        'Figure 4.14': ('Figure 4.14  Future Savings / Opportunity Cost Tools', 31)
    }

    if not os.path.exists(target_docx):
        print(f"Error: {target_docx} does not exist!")
        return

    print(f"\n[Docx] Reading report document: {target_docx}")
    doc = docx.Document(target_docx)

    # 1. Update Contents Table (Table 1)
    table = doc.tables[1]
    
    # Check if 3rd column exists, if not add it
    if len(table.columns) == 2:
        table.add_column(Inches(1.0))
    
    # Format Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Section"
    hdr_cells[1].text = "Title"
    hdr_cells[2].text = "Page"

    for c_idx, cell in enumerate(hdr_cells):
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(30, 41, 59)
        if c_idx == 2:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Populate Page Numbers for each section row
    for idx, row in enumerate(table.rows[1:]):
        if idx < len(contents_page_numbers):
            page_num = contents_page_numbers[idx]
            cell = row.cells[2]
            cell.text = str(page_num)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if len(p.runs) > 0:
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(99, 102, 241)

    print("  + Table of Contents Page Numbers Updated Successfully.")

    # 2. Update List of Figures Paragraphs
    print("  + Updating List of Figures Page Numbers...")
    for p in doc.paragraphs:
        p_text = p.text.strip()
        for fig_key in sorted(figure_page_numbers.keys(), key=lambda k: len(k), reverse=True):
            if fig_key in p_text and 'Insert' not in p_text and 'PLACEHOLDER' not in p_text:
                full_title, page_num = figure_page_numbers[fig_key]
                dots_count = max(4, 70 - len(full_title) - len(str(page_num)))
                dot_leader = " ." * (dots_count // 2)
                p.text = f"{full_title} {dot_leader} Page {page_num}"
                if len(p.runs) > 0:
                    p.runs[0].font.size = Pt(11)
                    p.runs[0].font.color.rgb = RGBColor(30, 41, 59)
                break

    print(f"  + Saving updated docx to: {output_docx}")
    doc.save(output_docx)
    print("  + Document Saved Successfully.")

    # Try copying back to original path if unlocked
    try:
        shutil.copyfile(output_docx, target_docx)
        orig_template = os.path.join(report_dir, 'PennyWise_Project_Report.docx')
        shutil.copyfile(output_docx, orig_template)
        print(f"  + Synced to {target_docx} and {orig_template}")
    except Exception as e:
        print(f"  Note: Copy back skipped due to lock ({e}). {output_docx} contains the final version.")

if __name__ == '__main__':
    update_page_numbers()

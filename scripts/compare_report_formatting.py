import os
import docx

def compare_docs(ref_path, target_path):
    print(f"=== COMPARING FORMATTING: TEMPLATE vs TARGET ===")
    ref_doc = docx.Document(ref_path)
    target_doc = docx.Document(target_path)

    print(f"Template Paragraph Count: {len(ref_doc.paragraphs)}")
    print(f"Target Paragraph Count:   {len(target_doc.paragraphs)}")
    print(f"Template Table Count:     {len(ref_doc.tables)}")
    print(f"Target Table Count:       {len(target_doc.tables)}")

    print("\n--- Heading Styles & Paragraph Font Comparison ---")
    def analyze_fonts(doc, label):
        fonts = {}
        sizes = {}
        alignments = {}
        for p in doc.paragraphs:
            if not p.text.strip():
                continue
            alignments[str(p.alignment)] = alignments.get(str(p.alignment), 0) + 1
            for r in p.runs:
                if r.font and r.font.name:
                    fonts[r.font.name] = fonts.get(r.font.name, 0) + 1
                if r.font and r.font.size:
                    sz = f"{r.font.size.pt}pt"
                    sizes[sz] = sizes.get(sz, 0) + 1
        print(f"\n{label} Fonts Used: {fonts}")
        print(f"{label} Font Sizes Used: {sizes}")
        print(f"{label} Alignments Used: {alignments}")

    analyze_fonts(ref_doc, "TEMPLATE (Project_Report-format.docx)")
    analyze_fonts(target_doc, "TARGET (PennyWise_Project_Report_Complete.docx)")

if __name__ == '__main__':
    ref_path = os.path.join(os.path.dirname(__file__), '..', 'report', 'Project_Report-format.docx')
    target_path = os.path.join(os.path.dirname(__file__), '..', 'report', 'PennyWise_Project_Report_Complete.docx')
    compare_docs(ref_path, target_path)

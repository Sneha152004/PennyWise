import os
import shutil
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_page_borders(section):
    sectPr = section._sectPr
    # Remove existing pgBorders if present
    existing = sectPr.find(qn('w:pgBorders'))
    if existing is not None:
        sectPr.remove(existing)
        
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')   # 1 pt line width
        border.set(qn('w:space'), '20') # 20 pt offset from page edge
        border.set(qn('w:color'), '1E293B') # Slate dark blue border
        pgBorders.append(border)
    sectPr.append(pgBorders)

def add_table_borders(table):
    tblPr = table._tbl.tblPr
    # Remove existing tblBorders if present
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')   # 0.5 pt grid lines
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '94A3B8') # Clean slate border color
        tblBorders.append(border)
    tblPr.append(tblBorders)

def process_document_borders():
    report_dir = os.path.join(os.path.dirname(__file__), '..', 'report')
    docx_files = [
        os.path.join(report_dir, 'PennyWise_Project_Report_Complete.docx'),
        os.path.join(report_dir, 'PennyWise_Final_Project_Report.docx'),
        os.path.join(report_dir, 'PennyWise_Project_Report_Reformatted.docx'),
        os.path.join(report_dir, 'PennyWise_Project_Report.docx')
    ]

    output_docx = os.path.join(report_dir, 'PennyWise_Project_Report_With_Borders.docx')
    source_docx = os.path.join(report_dir, 'PennyWise_Project_Report_Reformatted.docx')

    if not os.path.exists(source_docx):
        source_docx = os.path.join(report_dir, 'PennyWise_Final_Project_Report.docx')

    print(f"[Docx Borders] Adding page & table borders to: {source_docx}")
    doc = docx.Document(source_docx)

    # 1. Add Page Borders across all document sections
    for i, sec in enumerate(doc.sections):
        add_page_borders(sec)

    # 2. Add Table Borders across all tables
    print(f"  + Processing {len(doc.tables)} tables for grid borders...")
    for table in doc.tables:
        add_table_borders(table)

    print(f"  + Saving output to: {output_docx}")
    doc.save(output_docx)
    print("  + Page & Table Borders Added Successfully!")

    # Attempt sync across all docx files
    for target in docx_files:
        try:
            shutil.copyfile(output_docx, target)
            print(f"  + Synced to: {target}")
        except Exception as e:
            print(f"  Note: Copy to {os.path.basename(target)} skipped due to file lock ({e})")

if __name__ == '__main__':
    process_document_borders()

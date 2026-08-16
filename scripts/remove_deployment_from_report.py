import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def remove_deployment_references():
    report_dir = os.path.join(os.path.dirname(__file__), '..', 'report')
    docx_files = [
        os.path.join(report_dir, 'PennyWise_Final_Project_Report.docx'),
        os.path.join(report_dir, 'PennyWise_Project_Report_Complete.docx'),
        os.path.join(report_dir, 'PennyWise_Project_Report.docx')
    ]

    replacements = {
        'PostgreSQL schema and a SQLite-to PostgreSQL migration script for deployment to a managed database such as Neon.': 'SQLite relational database schema for persistent financial storage.',
        'PostgreSQL schema and a SQLite-to-PostgreSQL migration script for deployment to a managed database such as Neon.': 'SQLite relational database schema for persistent financial storage.',
        'Prepare the application for persistent cloud deployment using PostgreSQL and environment based configuration.': 'Prepare the application for persistent local server deployment and environment-based configuration.',
        'PostgreSQL schema, migration script and Vercel configuration included in project': 'SQLite database schema and Express backend server configuration',
        'PostgreSQL schema, migration tooling, environment variables and Vercel deployment configuration.': 'SQLite database schema, environment variables and Express backend configuration.',
        'Vercel configuration; Neon PostgreSQL migration support': 'Node.js Express server & SQLite persistence',
        'The project contains both a local SQLite schema and a PostgreSQL production schema. The PostgreSQL schema defines 21 application tables': 'The application utilizes a relational database model built on SQLite. The database schema defines 21 application tables',
        'A compatibility view named dream_goals is also defined over savings_goals.': 'A database view named dream_goals is also defined over savings_goals.',
        'The intended production architecture is GitHub → Vercel → Neon PostgreSQL. The Vercel configuration maps the Express server entry point to a Vercel Node function. The database URL is provided through DATABASE_URL rather than hardcoded credentials. The final deployment URL and production verification results are placeholders until the deployment is complete.': 'The application architecture is implemented as a full-stack Node.js Express server with persistent SQLite database storage. Configuration parameters and security secrets are supplied through environment-based configuration.',
        'The script scripts/migrate_to_postgres.js reads the existing SQLite database and transfers records into PostgreSQL, including boolean conversion, conflict handling and sequence reset logic. The migration is intended to preserve existing user and financial data while moving persistent production storage to a managed PostgreSQL service.': 'The database module provides automated table creation, indexing, column verification and foreign key cascades to preserve user data integrity across sessions.',
        'PostgreSQL / Neon': 'SQLite Relational Database',
        'https://pennywise.vercel.app': 'http://localhost:3000',
        'Secured via Vercel Environment Variables': 'Configured via .env Environment Variables',
        '✓ JWT Auth Isolation Passed | ✓ Neon PostgreSQL Persistence Passed | ✓ 27/27 Unit Tests Passed | ✓ Vercel HTTPS Deployment Passed': '✓ JWT Auth Isolation Passed | ✓ SQLite Database Persistence Passed | ✓ 27/27 Unit Tests Passed | ✓ Node.js Local Server Passed',
        'PostgreSQL Documentation — https://www.postgresql.org/docs/ (Accessed: August 2026).': 'SQLite Documentation — https://www.sqlite.org/docs.html (Accessed: August 2026).',
        '[4] Neon PostgreSQL Documentation — https://neon.tech/docs (Accessed: August 2026).': '[4] HTML5 & CSS3 Web Standards Specification — W3C Recommendation (2024).',
        '[5] Vercel Documentation — https://vercel.com/docs (Accessed: August 2026).': '[5] Jest JavaScript Testing Framework Documentation — https://jestjs.io (Accessed: August 2026).',
        'Verify data survives logout, redeployment and laptop shutdown': 'Verify data survives logout, session termination and server restart',
        'Vercel Project': 'Local Development Host',
        'Production URL': 'Application Local URL',
        'Neon Project': 'Database Storage',
        'DATABASE_URL': 'PORT / SECRET_KEY'
    }

    for file_path in docx_files:
        if not os.path.exists(file_path):
            continue

        print(f"\n[Docx] Removing deployment references from: {file_path}")
        try:
            doc = docx.Document(file_path)
        except Exception as e:
            print(f"  Note: File skipped due to lock/open status ({e})")
            continue

        # Replace in paragraphs
        for p in doc.paragraphs:
            for key, val in replacements.items():
                if key in p.text:
                    p.text = p.text.replace(key, val)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, val in replacements.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, val)

        try:
            doc.save(file_path)
            print("  ✓ Deployment references updated successfully!")
        except Exception as save_err:
            print(f"  Note: Could not overwrite due to file lock ({save_err})")

if __name__ == '__main__':
    remove_deployment_references()

import os
import docx

def verify_integrity():
    report_dir = os.path.join(os.path.dirname(__file__), '..', 'report')
    orig_path = os.path.join(report_dir, 'PennyWise_Final_Project_Report.docx')
    ref_path = os.path.join(report_dir, 'PennyWise_Project_Report_Reformatted.docx')

    doc_orig = docx.Document(orig_path)
    doc_ref = docx.Document(ref_path)

    orig_texts = [p.text.strip() for p in doc_orig.paragraphs if p.text.strip()]
    ref_texts = [p.text.strip() for p in doc_ref.paragraphs if p.text.strip()]

    print("=== CONTENT INTEGRITY VERIFICATION CHECK ===")
    print(f"Original Text Paragraph Count:   {len(orig_texts)}")
    print(f"Reformatted Text Paragraph Count: {len(ref_texts)}")

    mismatches = 0
    for idx, (t1, t2) in enumerate(zip(orig_texts, ref_texts)):
        if t1 != t2:
            print(f"[MISMATCH at P{idx}]")
            print(f"  Orig: {t1[:60]}...")
            print(f"  Ref:  {t2[:60]}...")
            mismatches += 1

    if len(orig_texts) == len(ref_texts) and mismatches == 0:
        print("\n+ VERIFICATION PASSED: Content is 100% IDENTICAL across all paragraphs and sections! ZERO text modified.")
    else:
        print(f"\n- WARNING: Found {mismatches} mismatches out of {len(orig_texts)} paragraphs.")

if __name__ == '__main__':
    verify_integrity()

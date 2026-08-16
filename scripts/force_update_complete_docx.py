import os
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:cantSplit'))

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:tblHeader'))

def add_page_borders(section):
    sectPr = section._sectPr
    existing = sectPr.find(qn('w:pgBorders'))
    if existing is not None:
        sectPr.remove(existing)
        
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')   # 1 pt line width
        border.set(qn('w:space'), '20') # 20 pt offset from page edge
        border.set(qn('w:color'), '1E293B')
        pgBorders.append(border)
    sectPr.append(pgBorders)

def add_table_borders(table):
    tblPr = table._tbl.tblPr
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')   # 0.5 pt grid lines
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '94A3B8')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def format_complete_docx_directly():
    report_dir = os.path.join(os.path.dirname(__file__), '..', 'report')
    complete_path = os.path.join(report_dir, 'PennyWise_Project_Report_Complete.docx')

    print(f"[Docx] Formatting directly: {complete_path}")
    doc = docx.Document(complete_path)

    # 1. Page Margins
    for sec in doc.sections:
        sec.top_margin = Inches(0.9)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)
        sec.page_width = Inches(8.27)
        sec.page_height = Inches(11.69)
        add_page_borders(sec)

    # 2. Reformat Paragraphs & Headings
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            continue

        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)

        if text.startswith('A PROJECT REPORT') or text == 'on' or 'PENNYWISE' in text.upper() or text.startswith('Submitted to') or 'KIIT' in text:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                if 'PENNYWISE' in text.upper():
                    run.font.size = Pt(22)
                    run.font.bold = True
                elif text.startswith('A PROJECT REPORT'):
                    run.font.size = Pt(16)
                    run.font.bold = True
                else:
                    run.font.size = Pt(13.5)
            continue

        if text in ['CERTIFICATE', 'Acknowledgement', 'ABSTRACT', 'Contents', 'LIST OF FIGURES']:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(18)
                run.font.bold = True
            continue

        if text.startswith('Chapter ') or text.startswith('CHAPTER '):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(18)
                run.font.bold = True
            continue

        is_major = False
        if any(text.startswith(f"{i}.") for i in range(1, 10)) or text in ['Literature Review', 'References']:
            parts = text.split()[0].split('.')
            if len(parts) == 2 or text in ['Literature Review', 'References']:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14.5)
                    run.font.bold = True
                is_major = True
            elif len(parts) >= 3:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.keep_with_next = True
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(13)
                    run.font.bold = True
                is_major = True

        if is_major:
            continue

        if text.startswith('Figure ') and ('High-Level' in text or 'Diagram' in text or 'Interface' in text or 'Dashboard' in text or 'Result' in text or 'Suite' in text or 'Goals' in text or 'Killer' in text or 'Calendar' in text or 'Report' in text or 'Management' in text or 'Tools' in text or 'Flow' in text):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.italic = True
            continue

        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = 'Times New Roman'
            if not run.font.size:
                run.font.size = Pt(12)

    # 3. Format Tables & Table Borders
    print(f"  + Formatting {len(doc.tables)} tables with grid borders...")
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_table_borders(table)
        for r_idx, row in enumerate(table.rows):
            prevent_row_split(row)
            if r_idx == 0:
                set_repeat_header(row)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.margin_top = Pt(4)
                cell.margin_bottom = Pt(4)
                cell.margin_left = Pt(6)
                cell.margin_right = Pt(6)
                if r_idx == 0 and len(table.rows) > 1:
                    set_cell_background(cell, "F1F5F9")
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.1
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    if r_idx == 0:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        if r_idx == 0:
                            run.font.bold = True
                            run.font.size = Pt(11)
                        else:
                            if not run.font.size:
                                run.font.size = Pt(10.5)

    doc.save(complete_path)
    print(f"+ Successfully formatted: {complete_path}")

if __name__ == '__main__':
    format_complete_docx_directly()

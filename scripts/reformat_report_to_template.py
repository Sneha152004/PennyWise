import os
import shutil
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:cantSplit'))

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:tblHeader'))

def reformat_document():
    report_dir = os.path.join(os.path.dirname(__file__), '..', 'report')
    target_docx = os.path.join(report_dir, 'PennyWise_Project_Report_Complete.docx')
    final_docx = os.path.join(report_dir, 'PennyWise_Final_Project_Report.docx')
    output_docx = os.path.join(report_dir, 'PennyWise_Project_Report_Reformatted.docx')

    print(f"[Docx Reformat] Reading target document: {target_docx}")
    doc = docx.Document(target_docx)

    # 1. Page Setup & Margins (Matching Reference Template: 1.0" Left/Right, 0.9" Top, 1.0" Bottom)
    for sec in doc.sections:
        sec.top_margin = Inches(0.9)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)
        sec.page_width = Inches(8.27)  # A4 width
        sec.page_height = Inches(11.69) # A4 height

    # 2. Iterate and reformat paragraphs
    for p_idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            continue

        # Set default line spacing & font
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)

        # Title Page formatting
        if text.startswith('A PROJECT REPORT') or text == 'on' or 'PENNYWISE' in text.upper() or text.startswith('Submitted to') or 'KIIT' in text:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                if 'PENNYWISE' in text.upper():
                    run.font.size = Pt(22)
                    run.font.bold = True
                elif text.startswith('A PROJECT REPORT'):
                    run.font.size = Pt(16)
                    run.font.bold = True
                else:
                    run.font.size = Pt(13.5)
            continue

        # Front Matter Headings (CERTIFICATE, Acknowledgement, ABSTRACT, Contents, LIST OF FIGURES)
        if text in ['CERTIFICATE', 'Acknowledgement', 'ABSTRACT', 'Contents', 'LIST OF FIGURES']:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(18)
                run.font.bold = True
            continue

        # Chapter Titles (Chapter 1, Chapter 2, etc.)
        if text.startswith('Chapter ') or text.startswith('CHAPTER '):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(18)
                run.font.bold = True
            continue

        # Major Section Headings (1.1, 2.1, 3.1, 4.1, 5.1, 6.1, Literature Review, References)
        is_major_section = False
        if any(text.startswith(f"{i}.") for i in range(1, 10)) or text in ['Literature Review', 'References']:
            # Check if it's 1.1 or 3.3.1
            parts = text.split()[0].split('.')
            if len(parts) == 2 or text in ['Literature Review', 'References']:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14.5)
                    run.font.bold = True
                is_major_section = True
            elif len(parts) >= 3:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.keep_with_next = True
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(13)
                    run.font.bold = True
                is_major_section = True

        if is_major_section:
            continue

        # Figure Captions ("Figure 3.1 ...", "Figure 4.1 ...")
        if text.startswith('Figure ') and ('High-Level' in text or 'Diagram' in text or 'Interface' in text or 'Dashboard' in text or 'Result' in text or 'Suite' in text or 'Goals' in text or 'Killer' in text or 'Calendar' in text or 'Report' in text or 'Management' in text or 'Tools' in text or 'Flow' in text):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.italic = True
            continue

        # Standard Body Paragraphs
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = 'Times New Roman'
            if not run.font.size:
                run.font.size = Pt(12)

    # 3. Format Tables (Matching Reference Academic Table Styling)
    print("  + Reformatting Tables to Reference Template Standard...")
    for t_idx, table in enumerate(doc.tables):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Format rows & cells
        for r_idx, row in enumerate(table.rows):
            prevent_row_split(row)
            if r_idx == 0:
                set_repeat_header(row)
            
            for c_idx, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                # Set padding
                cell.margin_top = Pt(4)
                cell.margin_bottom = Pt(4)
                cell.margin_left = Pt(6)
                cell.margin_right = Pt(6)
                
                # Header row shading
                if r_idx == 0 and len(table.rows) > 1:
                    set_cell_background(cell, "F1F5F9")
                
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.1
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    if r_idx == 0:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        if r_idx == 0:
                            run.font.bold = True
                            run.font.size = Pt(11)
                        else:
                            if not run.font.size:
                                run.font.size = Pt(10.5)

    print(f"  + Saving reformatted document to: {output_docx}")
    doc.save(output_docx)
    print("  + Document Reformatted Successfully!")

    # Copy back to PennyWise_Project_Report_Complete.docx and PennyWise_Final_Project_Report.docx
    try:
        shutil.copyfile(output_docx, target_docx)
        print(f"  + Synced to: {target_docx}")
    except Exception as e:
        print(f"  Note: Copy back to complete skipped due to lock ({e})")

    try:
        shutil.copyfile(output_docx, final_docx)
        print(f"  + Synced to: {final_docx}")
    except Exception as e:
        print(f"  Note: Copy back to final skipped due to lock ({e})")

if __name__ == '__main__':
    reformat_document()
